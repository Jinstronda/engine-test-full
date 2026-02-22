"""Document parser tool — extracts clean text from files and URLs.

Supported:
  - PDF  (.pdf)              via pymupdf (pip install pymupdf)
  - DOCX (.docx / .doc)     via python-docx (pip install python-docx)
  - Text (.txt .md .csv .json .html .htm)  built-in
  - URL  (http/https)        via httpx + basic HTML stripping

Install: pip install pymupdf python-docx httpx
"""

from __future__ import annotations

import logging
import os
import re
from typing import Optional

from langchain_core.tools import tool

from app.tools import register

logger = logging.getLogger(__name__)

_MAX_TEXT = 8_000  # chars returned to the agent

_EXTENSION_MAP = {
    ".pdf": "pdf",
    ".docx": "docx",
    ".doc": "docx",
    ".txt": "text",
    ".md": "text",
    ".csv": "text",
    ".json": "text",
    ".html": "text",
    ".htm": "text",
}


# ── Internal parsers ─────────────────────────────────────────────────────────


def _detect_source_type(source: str) -> str:
    if source.startswith("http://") or source.startswith("https://"):
        return "url"
    ext = os.path.splitext(source)[1].lower()
    return _EXTENSION_MAP.get(ext, "text")


def _strip_html(html: str) -> str:
    """Very lightweight HTML → plain text stripper (no extra deps)."""
    # Remove script / style blocks
    text = re.sub(r"<(script|style)[^>]*>.*?</\1>", "", html, flags=re.DOTALL | re.IGNORECASE)
    # Remove all tags
    text = re.sub(r"<[^>]+>", " ", text)
    # Collapse whitespace
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _parse_pdf(file_path: str) -> tuple[str, str, Optional[str]]:
    """Returns (title, text, error)."""
    try:
        import fitz  # pymupdf
    except ImportError:
        return "", "", "pymupdf is not installed. Run: pip install pymupdf"
    try:
        doc = fitz.open(file_path)
        pages_text = [page.get_text("text") for page in doc if page.get_text("text").strip()]
        full_text = "\n\n".join(pages_text)
        meta = doc.metadata or {}
        title = meta.get("title") or os.path.basename(file_path)
        doc.close()
        return title, full_text, None
    except Exception as e:
        return "", "", str(e)


def _parse_docx(file_path: str) -> tuple[str, str, Optional[str]]:
    """Returns (title, text, error)."""
    try:
        from docx import Document
    except ImportError:
        return "", "", "python-docx is not installed. Run: pip install python-docx"
    try:
        doc = Document(file_path)
        parts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                parts.append(" | ".join(c.text.strip() for c in row.cells))
        full_text = "\n\n".join(parts)
        title = (doc.core_properties.title if doc.core_properties else None) or os.path.basename(file_path)
        return title, full_text, None
    except Exception as e:
        return "", "", str(e)


def _parse_text_file(file_path: str) -> tuple[str, str, Optional[str]]:
    try:
        with open(file_path, "r", encoding="utf-8", errors="replace") as f:
            content = f.read()
        return os.path.basename(file_path), content, None
    except Exception as e:
        return "", "", str(e)


async def _parse_url(url: str) -> tuple[str, str, Optional[str]]:
    """Returns (title, text, error)."""
    try:
        import httpx
    except ImportError:
        return "", "", "httpx is not installed. Run: pip install httpx"
    try:
        async with httpx.AsyncClient(follow_redirects=True, timeout=15.0) as client:
            resp = await client.get(url, headers={"User-Agent": "fabriq-bot/1.0"})
            resp.raise_for_status()

        html = resp.text
        # Extract <title>
        title_m = re.search(r"<title[^>]*>(.*?)</title>", html, re.IGNORECASE | re.DOTALL)
        title = title_m.group(1).strip() if title_m else url
        text = _strip_html(html)
        return title, text, None
    except Exception as e:
        return "", "", str(e)


# ── Tool ─────────────────────────────────────────────────────────────────────


@register
@tool
async def parse_document(source: str) -> str:
    """Extract clean text from a document (PDF, DOCX, text file) or a web URL.

    Useful for reading resumes, portfolios, job descriptions, or any document
    that an agent needs to analyse.

    Args:
        source: An absolute file path (e.g. "/tmp/resume.pdf") or a URL
            (e.g. "https://example.com/about"). File type is auto-detected
            from the extension.  Supported: .pdf, .docx, .doc, .txt, .md,
            .csv, .json, .html, and any http/https URL.

    Returns:
        The document title followed by up to 8,000 characters of extracted
        plain text. Returns an error message on failure.
    """
    source_type = _detect_source_type(source)

    if source_type == "url":
        title, text, error = await _parse_url(source)
    elif source_type == "pdf":
        title, text, error = _parse_pdf(source)
    elif source_type == "docx":
        title, text, error = _parse_docx(source)
    else:
        title, text, error = _parse_text_file(source)

    if error:
        return f"Failed to parse '{source}': {error}"

    if not text.strip():
        return f"No text content found in '{source}'."

    text_snippet = text[:_MAX_TEXT] + ("..." if len(text) > _MAX_TEXT else "")
    return f"[{title}]\n\n{text_snippet}"
